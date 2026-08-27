"""Generate lyrics document in DOCX format."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from splitaudio.metadata import TrackMeta

log = logging.getLogger(__name__)

_SECTION_CN = {
    "verse": "[verse] · 主歌",
    "chorus": "[chorus] · 副歌",
    "bridge": "[bridge] · 桥段",
    "intro": "[intro] · 前奏",
    "outro": "[outro] · 尾奏",
}


def write_lyrics_docx(meta: TrackMeta, dst: Path) -> None:
    """Generate a DOCX file with the lyrics from TrackMeta."""
    dst.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Setup Normal style with Chinese font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = 1  # CENTER
    title_run = title_para.add_run(meta.title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Subtitle: artist · duration
    dur_m = int(meta.duration // 60)
    dur_s = int(meta.duration % 60)
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = 1  # CENTER
    sub_text = f"{meta.artist} · {dur_m:02d}:{dur_s:02d}"
    sub_run = subtitle_para.add_run(sub_text)
    sub_run.font.size = Pt(11)
    sub_run.font.name = "Times New Roman"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Spacer
    doc.add_paragraph()

    # Lyrics
    if not meta.lyrics.strip():
        empty_para = doc.add_paragraph("（未找到内嵌歌词）")
        empty_para.alignment = 1
    else:
        for section in meta.sections:
            if section.marker:
                marker_lower = section.marker.strip("[]").lower()
                display = _SECTION_CN.get(marker_lower, section.marker)
                sec_para = doc.add_paragraph()
                sec_para.paragraph_format.space_before = Pt(12)
                sec_para.paragraph_format.space_after = Pt(6)
                sec_run = sec_para.add_run(display)
                sec_run.bold = True
                sec_run.font.size = Pt(12)
                sec_run.font.name = "Times New Roman"
                sec_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

            for line in section.lines:
                line_para = doc.add_paragraph()
                line_para.paragraph_format.line_spacing = 1.5
                line_run = line_para.add_run(line)
                line_run.font.size = Pt(12)
                line_run.font.name = "Times New Roman"
                line_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(str(dst))
    log.info("歌词文档已生成: %s → %s", meta.title, dst.name)
